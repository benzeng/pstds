# pstds/export/docx_exporter.py
# Word 导出器 - Phase 6 Task 2 (P6-T2)

from typing import Dict, Any
from datetime import date, datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.shared import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = None
    Inches = None
    Pt = None
    RGBColor = None
    WD_ALIGN_PARAGRAPH = None
    OxmlElement = None


class WordExporter:
    """
    Word 导出器

    使用 python-docx 导出为 Word 文档，含封面和免责声明。
    """

    def export_backtest_result(self, result: Dict[str, Any]) -> bool:
        """
        导出回测结果为 Word 文档

        Args:
            result: 回测结果字典

        Returns:
            是否导出成功
        """
        if not DOCX_AVAILABLE:
            print("python-docx 未安装，无法导出 Word 文档")
            return False

        try:
            doc = Document()

            # 封面
            self._add_cover_page(doc, result)

            # 换页
            doc.add_page_break()

            # 目录
            doc.add_heading("目录", level=1)
            doc.add_paragraph("1. 绩效摘要")
            doc.add_paragraph("2. 净值曲线")
            doc.add_paragraph("3. 交易统计")
            doc.add_paragraph("4. 每日快照")
            doc.add_page_break()

            # 绩效摘要
            self._add_performance_section(doc, result)

            # 交易统计
            doc.add_page_break()
            self._add_trades_section(doc, result)

            # 免责声明
            doc.add_page_break()
            self._add_disclaimer(doc)

            # 保存文件
            filename = f"backtest_{result.get('symbol', 'N/A')}_{datetime.now().strftime('%Y%m%d')}.docx"
            doc.save(filename)
            print(f"Word 文档已保存: {filename}")
            return True

        except Exception as e:
            print(f"导出 Word 文档失败: {e}")
            return False

    def export_analysis(self, decision: Dict[str, Any]) -> bool:
        """
        导出分析结果为 Word 文档

        Args:
            decision: 分析决策字典

        Returns:
            是否导出成功
        """
        if not DOCX_AVAILABLE:
            print("python-docx 未安装，无法导出 Word 文档")
            return False

        try:
            doc = Document()

            # 封面
            self._add_cover_page(doc, {"title": "分析报告", "symbol": decision.get('symbol', 'N/A')})

            # 目录
            doc.add_heading("目录", level=1)
            doc.add_paragraph("1. 决策信息")
            doc.add_paragraph("2. 价格目标")
            doc.add_paragraph("3. 风险因素")
            doc.add_paragraph("4. 数据来源")
            doc.add_page_break()

            # 决策信息
            self._add_decision_section(doc, decision)

            # 价格目标
            doc.add_page_break()
            self._add_target_price_section(doc, decision)

            # 风险因素
            doc.add_page_break()
            self._add_risk_factors_section(doc, decision)

            # 数据来源
            doc.add_page_break()
            self._add_data_sources_section(doc, decision)

            # 免责声明
            doc.add_page_break()
            self._add_disclaimer(doc)

            # 保存文件
            filename = f"analysis_{decision.get('symbol', 'N/A')}_{datetime.now().strftime('%Y%m%d')}.docx"
            doc.save(filename)
            print(f"Word 文档已保存: {filename}")
            return True

        except Exception as e:
            print(f"导出 Word 文档失败: {e}")
            return False

    def _add_cover_page(self, doc: Document, info: Dict[str, Any]):
        """添加封面"""
        # 标题
        title = info.get("title", "PSTDS 报告")
        heading = doc.add_heading(title, level=0)

        # 副标题
        doc.add_paragraph(f"股票代码: {info.get('symbol', 'N/A')}")
        doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y年%m月%d日 %H:%M')}")

        # 分割线
        paragraph = doc.add_paragraph()
        paragraph.add_run().add_break()
        p = OxmlElement('w:pPr')
        p.set(qnval('600000000'))

    def _add_performance_section(self, doc: Document, result: Dict[str, Any]):
        """添加绩效摘要部分"""
        doc.add_heading("1. 绩效摘要", level=2)

        # 基本指标
        doc.add_paragraph("**初始资金**:")
        doc.add_paragraph(f"${result.get('initial_capital', 0):.2f}")

        doc.add_paragraph("**最终净值**:")
        doc.add_paragraph(f"${result.get('final_nav', 0):.2f}")

        # 绩效指标表格
        table = doc.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '指标'
        hdr_cells[1].text = '数值'

        metrics = [
            ("总收益率", f"{result.get('total_return', 0):.2f}%"),
            ("年化收益率", f"{result.get('annualized_return', 0):.2f}%"),
            ("最大回撤", f"{result.get('max_drawdown', 0):.2f}%"),
            ("夏普比率", f"{result.get('sharpe_ratio', 0):.2f}"),
            ("卡尔马比率", f"{result.get('calmar_ratio', 0):.2f}"),
        ]

        for metric_name, metric_value in metrics:
            row_cells = table.add_row().cells
            row_cells[0].text = metric_name
            row_cells[1].text = metric_value

    def _add_decision_section(self, doc: Document, decision: Dict[str, Any]):
        """添加决策信息部分"""
        doc.add_heading("1. 决策信息", level=2)

        doc.add_paragraph("**决策**:")
        doc.add_paragraph(decision.get('action', 'N/A'))

        doc.add_paragraph("**置信度**:")
        doc.add_paragraph(f"{decision.get('confidence', 0):.2f}")

        doc.add_paragraph("**信心度**:")
        doc.add_paragraph(decision.get('conviction', 'N/A'))

        doc.add_paragraph("**主要理由**:")
        doc.add_paragraph(decision.get('primary_reason', 'N/A'))

    def _add_target_price_section(self, doc: Document, decision: Dict[str, Any]):
        """添加价格目标部分"""
        doc.add_heading("2. 价格目标", level=2)

        if decision.get('target_price_low'):
            doc.add_paragraph("**目标价下限**:")
            doc.add_paragraph(f"${decision.get('target_price_low', 0):.2f}")

        if decision.get('target_price_high'):
            doc.add_paragraph("**目标价上限**:")
            doc.add_paragraph(f"${decision.get('target_price_high', 0):.2f}")

        doc.add_paragraph("**时间框架**:")
        doc.add_paragraph(decision.get('time_horizon', 'N/A'))

    def _add_risk_factors_section(self, doc: Document, decision: Dict[str, Any]):
        """添加风险因素部分"""
        doc.add_heading("3. 风险因素", level=2)

        for risk in decision.get('risk_factors', []):
            doc.add_paragraph(f"- {risk}", style='List Bullet')

    def _add_data_sources_section(self, doc: Document, decision: Dict[str, Any]):
        """添加数据来源部分"""
        doc.add_heading("4. 数据来源", level=2)

        for source in decision.get('data_sources', []):
            doc.add_paragraph(f"- {source.get('name', 'N/A')}")

    def _add_disclaimer(self, doc: Document):
        """添加免责声明"""
        doc.add_page_break()
        doc.add_heading("免责声明", level=1)

        disclaimer_text = """
本报告由 PSTDS（个人专用股票交易决策系统）自动生成。

重要提示：
1. 投资有风险，入市须谨慎
2. 本系统不构成任何形式的投资建议
3. 开发者对投资损失不承担任何责任
4. 请在充分理解风险的前提下使用本系统
5. 本报告的任何内容仅供研究参考，不作为投资决策的唯一依据

© 2026 PSTDS - 个人专用股票交易决策系统
        """

        doc.add_paragraph(disclaimer_text)
